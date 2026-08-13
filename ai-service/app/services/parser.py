"""회의록 문서 파서 (TXT / PDF / DOCX)."""

import io
import logging
from pathlib import Path

logger = logging.getLogger(__name__)

SUPPORTED_EXTENSIONS = {".txt", ".pdf", ".docx"}
MAX_TRANSCRIPT_CHARS = 60_000


class UnsupportedFileTypeError(ValueError):
    """지원하지 않는 확장자."""


class EmptyDocumentError(ValueError):
    """텍스트를 추출하지 못한 문서."""


def parse_document(filename: str, content: bytes) -> str:
    """파일 확장자에 따라 적절한 파서로 텍스트를 추출한다."""

    extension = Path(filename or "").suffix.lower()
    if extension not in SUPPORTED_EXTENSIONS:
        raise UnsupportedFileTypeError(
            f"지원하지 않는 파일 형식입니다: {extension or '(확장자 없음)'} "
            f"(지원 형식: {', '.join(sorted(SUPPORTED_EXTENSIONS))})"
        )

    if extension == ".txt":
        text = _parse_txt(content)
    elif extension == ".pdf":
        text = _parse_pdf(content)
    else:
        text = _parse_docx(content)

    text = normalize(text)
    if not text:
        raise EmptyDocumentError("문서에서 텍스트를 추출하지 못했습니다. (스캔 이미지 PDF 등)")
    return text


def _parse_txt(content: bytes) -> str:
    for encoding in ("utf-8", "cp949", "euc-kr", "utf-16"):
        try:
            return content.decode(encoding)
        except UnicodeDecodeError:
            continue
    return content.decode("utf-8", errors="ignore")


def _parse_pdf(content: bytes) -> str:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(content))
    pages = [page.extract_text() or "" for page in reader.pages]
    logger.debug("PDF 파싱 완료: %d 페이지", len(pages))
    return "\n".join(pages)


def _parse_docx(content: bytes) -> str:
    import docx

    document = docx.Document(io.BytesIO(content))
    blocks = [paragraph.text for paragraph in document.paragraphs]

    # 표 안에 회의록이 정리된 경우도 흔하므로 표 셀도 함께 추출한다.
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                blocks.append(" | ".join(cells))

    return "\n".join(blocks)


def normalize(text: str) -> str:
    """공백 정리 + 과도하게 긴 회의록 절단(토큰 비용 방어)."""

    lines = [line.strip() for line in (text or "").splitlines()]
    cleaned = "\n".join(line for line in lines if line)
    if len(cleaned) > MAX_TRANSCRIPT_CHARS:
        logger.warning("회의록이 너무 길어 %d자로 절단합니다.", MAX_TRANSCRIPT_CHARS)
        cleaned = cleaned[:MAX_TRANSCRIPT_CHARS]
    return cleaned
